"""TASK-18 controlled proposal composition and local rendering."""

from __future__ import annotations

import html
import json
from datetime import datetime
from enum import StrEnum
from hashlib import sha256
from pathlib import Path
from typing import Self
from uuid import uuid4

from pydantic import BaseModel, ConfigDict, Field, model_validator


class ProposalApplicability(StrEnum):
    PROPOSAL_REQUIRED = "PROPOSAL_REQUIRED"
    CUSTOMER_FORM_ONLY = "CUSTOMER_FORM_ONLY"
    NO_PROPOSAL_REQUIRED = "NO_PROPOSAL_REQUIRED"
    NOT_ASSESSED = "NOT_ASSESSED"


class ProposalLifecycle(StrEnum):
    DRAFT = "DRAFT"
    ISSUED_FOR_REVIEW = "ISSUED_FOR_REVIEW"
    WITHDRAWN = "WITHDRAWN"
    SUPERSEDED = "SUPERSEDED"


class SectionRole(StrEnum):
    COVER = "COVER"
    EXECUTIVE_SUMMARY = "EXECUTIVE_SUMMARY"
    COMPLIANCE_RESPONSE = "COMPLIANCE_RESPONSE"
    SCOPE = "SCOPE"
    DELIVERABLES = "DELIVERABLES"
    SCHEDULE = "SCHEDULE"
    PRICING = "PRICING"
    PAYMENT_TERMS = "PAYMENT_TERMS"
    WARRANTY = "WARRANTY"
    VALIDITY = "VALIDITY"
    ASSUMPTIONS = "ASSUMPTIONS"
    EXCLUSIONS = "EXCLUSIONS"
    QUALIFICATIONS = "QUALIFICATIONS"
    DEVIATIONS = "DEVIATIONS"
    OPTIONS_ALTERNATES = "OPTIONS_ALTERNATES"
    ATTACHMENTS = "ATTACHMENTS"
    CUSTOM = "CUSTOM"


class ContentOrigin(StrEnum):
    STRUCTURED_SOURCE = "STRUCTURED_SOURCE"
    APPROVED_CONTENT_BLOCK = "APPROVED_CONTENT_BLOCK"
    OPERATOR_AUTHORED = "OPERATOR_AUTHORED"
    CALCULATED_PRESENTATION = "CALCULATED_PRESENTATION"


class ProposalProfile(BaseModel):
    model_config = ConfigDict(extra="forbid")
    profile_id: str = Field(default_factory=lambda: f"PPR-{uuid4().hex}")
    code: str
    name: str
    effective_from: datetime
    effective_until: datetime | None = None
    required_sections: tuple[SectionRole, ...]
    published: bool = False
    created_by: str
    created_at: datetime


class ProposalFamily(BaseModel):
    model_config = ConfigDict(extra="forbid")
    family_id: str = Field(default_factory=lambda: f"PFA-{uuid4().hex}")
    bid_id: str
    code: str
    applicability: ProposalApplicability
    title: str
    owner: str
    created_by: str
    created_at: datetime


class ProposalSection(BaseModel):
    model_config = ConfigDict(extra="forbid")
    section_id: str = Field(default_factory=lambda: f"PSE-{uuid4().hex}")
    role: SectionRole
    heading: str
    text: str
    origin: ContentOrigin
    source_ids: tuple[str, ...] = ()
    customer_visible: bool = True

    @model_validator(mode="after")
    def protect_internal(self) -> Self:
        lowered = self.text.casefold()
        blocked = ("margin", "gross profit", "internal approver", "mandate limit", "walk-away")
        if self.customer_visible and any(token in lowered for token in blocked):
            raise ValueError("internal-only content cannot enter customer output")
        if self.origin == ContentOrigin.STRUCTURED_SOURCE and not self.source_ids:
            raise ValueError("structured proposal content requires exact source IDs")
        return self


class ProposalVersion(BaseModel):
    model_config = ConfigDict(extra="forbid")
    proposal_version_id: str = Field(default_factory=lambda: f"PRV-{uuid4().hex}")
    family_id: str
    bid_id: str
    version_number: int = Field(ge=1)
    lifecycle: ProposalLifecycle = ProposalLifecycle.DRAFT
    profile_id: str
    presentation_currency: str
    sections: tuple[ProposalSection, ...]
    source_manifest: tuple[dict[str, str], ...] = ()
    commercial_baseline_id: str | None = None
    negotiated_position_id: str | None = None
    created_by: str
    created_at: datetime
    fingerprint: str | None = None

    @model_validator(mode="after")
    def freeze(self) -> Self:
        payload = json.dumps(self.model_dump(mode="json", exclude={"fingerprint"}), sort_keys=True)
        expected = sha256(payload.encode()).hexdigest()
        if self.fingerprint and self.fingerprint != expected:
            raise ValueError("proposal fingerprint mismatch")
        object.__setattr__(self, "fingerprint", expected)
        return self


class ProposalReview(BaseModel):
    model_config = ConfigDict(extra="forbid", frozen=True)
    review_id: str = Field(default_factory=lambda: f"PRR-{uuid4().hex}")
    proposal_version_id: str
    reviewer: str
    decision: str
    rationale: str
    reviewed_at: datetime


class RenderArtifact(BaseModel):
    model_config = ConfigDict(frozen=True)
    artifact_id: str = Field(default_factory=lambda: f"ART-{uuid4().hex}")
    proposal_version_id: str
    format: str
    relative_path: str
    media_type: str
    byte_size: int
    sha256: str
    verified: bool = True


def render_html(version: ProposalVersion) -> bytes:
    rows = "".join(
        f"<section><h2>{html.escape(section.heading)}</h2><p>{html.escape(section.text)}</p></section>"
        for section in version.sections
        if section.customer_visible
    )
    return (
        "<!doctype html><html><head><meta charset='utf-8'><title>Proposal</title>"
        f"</head><body>{rows}</body></html>"
    ).encode()


def render_json(version: ProposalVersion) -> bytes:
    return json.dumps(
        version.model_dump(mode="json"), sort_keys=True, separators=(",", ":")
    ).encode()


def render_docx(version: ProposalVersion, path: Path) -> None:
    from docx import Document

    document = Document()
    for section in version.sections:
        if section.customer_visible:
            document.add_heading(section.heading, level=1)
            document.add_paragraph(section.text)
    document.save(path)


def render_pdf(version: ProposalVersion, path: Path) -> None:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen.canvas import Canvas

    canvas = Canvas(str(path), pagesize=letter)
    y = 750
    for section in version.sections:
        if section.customer_visible:
            canvas.drawString(50, y, section.heading[:100])
            y -= 18
            canvas.drawString(60, y, section.text[:110])
            y -= 30
            if y < 80:
                canvas.showPage()
                y = 750
    canvas.save()


def write_artifacts(version: ProposalVersion, root: Path) -> tuple[RenderArtifact, ...]:
    root.mkdir(parents=True, exist_ok=True)
    outputs: list[tuple[str, bytes, str]] = [
        ("proposal.html", render_html(version), "text/html"),
        ("proposal.json", render_json(version), "application/json"),
    ]
    docx_path = root / "proposal.docx"
    render_docx(version, docx_path)
    outputs.append(
        (
            docx_path.name,
            docx_path.read_bytes(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    )
    pdf_path = root / "proposal.pdf"
    render_pdf(version, pdf_path)
    outputs.append((pdf_path.name, pdf_path.read_bytes(), "application/pdf"))
    artifacts: list[RenderArtifact] = []
    for name, data, media in outputs:
        target = root / name
        target.write_bytes(data)
        artifacts.append(
            RenderArtifact(
                proposal_version_id=version.proposal_version_id,
                format=target.suffix[1:].upper(),
                relative_path=name,
                media_type=media,
                byte_size=len(data),
                sha256=sha256(data).hexdigest(),
            )
        )
    return tuple(artifacts)
