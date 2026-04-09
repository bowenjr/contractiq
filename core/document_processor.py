"""
Document processor for ContractIQ.
Handles PDF and DOCX extraction, preserving structure.
"""

import re
from pathlib import Path
from typing import Dict, Any, List


class DocumentProcessor:

    def process(self, file_path: Path) -> Dict[str, Any]:
        """Extract text and metadata from a document file."""
        suffix = file_path.suffix.lower()
        if suffix == ".pdf":
            return self._process_pdf(file_path)
        elif suffix in (".docx", ".doc"):
            return self._process_docx(file_path)
        elif suffix == ".txt":
            return self._process_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

    def _process_pdf(self, file_path: Path) -> Dict[str, Any]:
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise ImportError("PyMuPDF not installed. Run: pip install pymupdf")

        doc = fitz.open(str(file_path))
        pages_content = []
        total_words = 0
        scanned_pages = 0

        for page_num, page in enumerate(doc):
            page_content = []

            # Method 1: Extract text blocks with position info
            blocks = page.get_text("dict")["blocks"]

            text_blocks = []
            for block in blocks:
                if block["type"] == 0:  # text block
                    block_text = ""
                    for line in block.get("lines", []):
                        line_text = ""
                        for span in line.get("spans", []):
                            line_text += span.get("text", "")
                        if line_text.strip():
                            block_text += line_text + "\n"
                    if block_text.strip():
                        text_blocks.append({
                            "text": block_text.strip(),
                            "bbox": block["bbox"],
                        })

            # Sort by vertical position then horizontal (handles multi-column)
            text_blocks.sort(
                key=lambda b: (round(b["bbox"][1] / 20) * 20, b["bbox"][0])
            )

            for block in text_blocks:
                page_content.append(block["text"])

            # Method 2: Extract tables if present
            try:
                tables = page.find_tables()
                if tables and tables.tables:
                    for table in tables.tables:
                        try:
                            df = table.to_pandas()
                            if not df.empty:
                                rows = []
                                for _, row in df.iterrows():
                                    cells = [
                                        str(v).strip()
                                        for v in row.values
                                        if str(v).strip() and str(v) != "nan"
                                    ]
                                    if cells:
                                        rows.append(" | ".join(cells))
                                if rows:
                                    page_content.append("\n".join(rows))
                        except Exception:
                            pass
            except Exception:
                pass

            page_text = "\n".join(page_content)
            word_count = len(page_text.split())
            if word_count < 10:
                scanned_pages += 1
            else:
                total_words += word_count
                pages_content.append(page_text)

        doc.close()

        if scanned_pages > max(1, len(pages_content)) * 0.5:
            print(
                f"  WARNING: {scanned_pages} pages appear to be scanned images. "
                f"Text extraction may be incomplete. Consider using a text-based PDF."
            )

        full_text = "\n\n".join(pages_content)
        full_text = self._clean_text(full_text)

        print(
            f"  PDF extraction: {total_words:,} words, "
            f"{len(full_text):,} chars, "
            f"{scanned_pages} scanned pages detected"
        )

        return {
            "text": full_text,
            "word_count": total_words,
            "page_count": len(pages_content),
            "doc_type": "PDF",
            "scanned_pages": scanned_pages,
        }

    def _process_docx(self, file_path: Path) -> Dict[str, Any]:
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")

        WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

        def get_text_from_element(element):
            texts = []
            for t in element.iter(f"{{{WORD_NS}}}t"):
                if t.text:
                    texts.append(t.text)
            return "".join(texts)

        def extract_table(tbl_element):
            rows = []
            for tr in tbl_element.findall(f".//{{{WORD_NS}}}tr"):
                cells = []
                for tc in tr.findall(f".//{{{WORD_NS}}}tc"):
                    cell_text = get_text_from_element(tc).strip()
                    if cell_text:
                        cells.append(cell_text)
                if cells:
                    rows.append(" | ".join(cells))
            return "\n".join(rows) if rows else ""

        def extract_shapes(doc):
            shape_texts = []
            for elem in doc.element.body.iter():
                tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
                if tag in ("txbxContent", "wsp", "txPr"):
                    text = get_text_from_element(elem).strip()
                    if text and len(text) > 5:
                        shape_texts.append(f"[TEXT BOX: {text}]")
            return shape_texts

        doc = Document(str(file_path))
        content_blocks = []

        # BLOCK 1: Body content in document order
        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            if tag == "p":
                text = get_text_from_element(element).strip()
                if text:
                    content_blocks.append(text)

            elif tag == "tbl":
                table_text = extract_table(element)
                if table_text:
                    content_blocks.append(table_text)

            elif tag == "sdt":
                # Structured document tags (form fields, content controls)
                text = get_text_from_element(element).strip()
                if text:
                    content_blocks.append(text)

        # BLOCK 2: Text boxes and shapes
        shapes = extract_shapes(doc)
        if shapes:
            content_blocks.append("\n--- TEXT BOXES ---")
            content_blocks.extend(shapes)

        # BLOCK 3: Headers and footers
        header_content = []
        for section in doc.sections:
            for hdr_ftr in [
                section.header, section.footer,
                section.even_page_header,
                section.first_page_header,
            ]:
                try:
                    if hdr_ftr and not hdr_ftr.is_linked_to_previous:
                        for para in hdr_ftr.paragraphs:
                            text = para.text.strip()
                            if text and len(text) > 3:
                                header_content.append(text)
                        for table in hdr_ftr.tables:
                            table_text = extract_table(table._tbl)
                            if table_text:
                                header_content.append(table_text)
                except Exception:
                    pass

        if header_content:
            content_blocks.append("\n--- DOCUMENT HEADERS/FOOTERS ---")
            content_blocks.extend(list(dict.fromkeys(header_content)))

        # BLOCK 4: Footnotes
        try:
            footnote_part = doc.part.footnotes_part
            if footnote_part:
                for fn in footnote_part._element.findall(f".//{{{WORD_NS}}}footnote"):
                    text = get_text_from_element(fn).strip()
                    if text and len(text) > 10:
                        content_blocks.append(f"[FOOTNOTE: {text}]")
        except Exception:
            pass

        full_text = "\n\n".join(b for b in content_blocks if b.strip())
        full_text = self._clean_text(full_text)
        word_count = len(full_text.split())

        print(
            f"  DOCX extraction: {word_count:,} words, "
            f"{len(full_text):,} chars from "
            f"{len(content_blocks)} blocks"
        )

        return {
            "text": full_text,
            "word_count": word_count,
            "page_count": max(1, word_count // 500),
            "doc_type": "DOCX",
        }

    def _process_txt(self, file_path: Path) -> Dict[str, Any]:
        text = file_path.read_text(encoding="utf-8", errors="replace")
        text = self._clean_text(text)
        word_count = len(text.split())
        return {
            "text": text,
            "word_count": word_count,
            "page_count": max(1, word_count // 500),
            "doc_type": "TXT",
        }

    def _clean_text(self, text: str) -> str:
        """Clean extracted text: normalize whitespace, remove junk."""
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = re.sub(r"[^\x09\x0A\x20-\x7E\u00A0-\uFFFF]", "", text)
        lines = []
        for line in text.split("\n"):
            line = re.sub(r" {3,}", "  ", line)
            lines.append(line)
        return "\n".join(lines).strip()

    def split_by_sections(self, text: str) -> List[Dict[str, str]]:
        """
        Split a legal document into sections by detecting heading patterns.

        Recognises:
          Article 1 / ARTICLE 1
          Section 1 / SECTION 1
          1.  Heading  /  1.1  Heading
          Schedule A / SCHEDULE A
          Appendix A / APPENDIX A  /  Exhibit A / EXHIBIT A

        Returns [{"heading": str, "content": str}].
        Falls back to chunk_text(max_chars=12000, overlap=800) when fewer
        than 3 headings are found.
        """
        heading_re = re.compile(
            r"^(?:"
            r"(?:ARTICLE|Article)\s+\d+"
            r"|(?:SECTION|Section)\s+\d+"
            r"|\d+\.\d+\s+[A-Z][A-Za-z]"   # 1.1 Definitions
            r"|\d+\.\s+[A-Z][A-Za-z]"       # 1. Payment Terms
            r"|(?:SCHEDULE|Schedule)\s+[A-Z0-9]"
            r"|(?:APPENDIX|Appendix)\s+[A-Z0-9]"
            r"|(?:EXHIBIT|Exhibit)\s+[A-Z0-9]"
            r")",
        )

        lines = text.split("\n")
        heading_positions: List[tuple] = []
        for i, line in enumerate(lines):
            stripped = line.strip()
            if stripped and len(stripped) < 120 and heading_re.match(stripped):
                heading_positions.append((i, stripped))

        if len(heading_positions) < 3:
            # Not enough structure — fall back to character chunking
            chunks = self.chunk_text(text, max_chars=12000, overlap=800)
            return [{"heading": f"Chunk {i + 1}", "content": c}
                    for i, c in enumerate(chunks)]

        sections: List[Dict[str, str]] = []

        # Preamble: everything before the first heading
        if heading_positions[0][0] > 0:
            preamble = "\n".join(lines[: heading_positions[0][0]]).strip()
            if preamble:
                sections.append({"heading": "Preamble", "content": preamble})

        for idx, (line_num, heading) in enumerate(heading_positions):
            next_line = (
                heading_positions[idx + 1][0]
                if idx + 1 < len(heading_positions)
                else len(lines)
            )
            content = "\n".join(lines[line_num:next_line]).strip()
            if content:
                sections.append({"heading": heading, "content": content})

        return sections

    def chunk_text(self, text: str, max_chars: int = 6000, overlap: int = 500):
        """Split text into overlapping chunks for LLM processing."""
        if len(text) <= max_chars:
            return [text]

        chunks = []
        start = 0
        while start < len(text):
            end = start + max_chars
            if end < len(text):
                boundary = text.rfind(". ", start, end)
                if boundary > start + max_chars // 2:
                    end = boundary + 1
            chunks.append(text[start:end])
            start = end - overlap
        return chunks
